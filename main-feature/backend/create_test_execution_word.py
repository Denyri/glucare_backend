import os
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

def create_word():
    doc = Document()
    
    # Orientasi landscape
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = 1
    section.page_width = new_width
    section.page_height = new_height

    doc.add_heading('Laporan Eksekusi Pengujian (Test Execution)', level=1)
    
    headers = [
        "No", "ID Test Case", "Fitur / Modul", "Pre-condition", "Skenario Pengujian", 
        "Data Uji", "Expected Result", "Actual Result", "Status", 
        "Komentar / Catatan", "Nama Tester", "Tanggal Test"
    ]
    
    data = [
        ["1", "TC-001", "Registrasi", "Pengguna belum memiliki akun terdaftar", "1. Buka halaman Registrasi\n2. Isi seluruh data\n3. Klik Daftar", "Email: andi@gmail.com\nPassword: user123", "Akun berhasil dibuat dan pengguna diarahkan ke halaman Login.", "Sistem berhasil menyimpan data pengguna baru ke basis data dan mengarahkan pengguna ke halaman Login setelah proses registrasi selesai.", "Pass", "-", "Deny Riansyah", "30 Juli 2026"],
        ["2", "TC-002", "Login", "Pengguna telah memiliki akun", "1. Buka halaman Login\n2. Masukkan email dan password\n3. Klik Masuk", "Email: user@gmail.com\nPassword: user123", "Pengguna berhasil login dan diarahkan ke Dashboard.", "Sistem berhasil melakukan autentikasi menggunakan email dan password yang valid, kemudian menampilkan halaman Dashboard.", "Pass", "-", "Deny Riansyah", "30 Juli 2026"],
        ["3", "TC-003", "Lengkapi Profil", "Pengguna telah login", "1. Membuka menu Lengkapi Profil\n2. Mengisi seluruh data\n3. Klik Simpan", "Data profil lengkap", "Data profil berhasil disimpan.", "Sistem berhasil menyimpan seluruh informasi profil pengguna dan menampilkan data yang telah diperbarui.", "Pass", "-", "Deny Riansyah", "30 Juli 2026"],
        ["4", "TC-004", "Analisis Klinis", "Profil telah lengkap", "1. Membuka menu Analisis Klinis\n2. Mengisi data klinis\n3. Klik Analisis", "Data klinis lengkap", "Sistem menampilkan hasil analisis risiko prediabetes.", "Sistem berhasil mengirim data klinis ke backend, menerima hasil dari layanan AI, dan menampilkan hasil analisis risiko.", "Pass", "-", "Deny Riansyah", "30 Juli 2026"],
        ["5", "TC-005", "Analisis Kuesioner", "Profil telah lengkap", "1. Membuka menu Analisis Kuesioner\n2. Menjawab pertanyaan\n3. Klik Analisis", "Jawaban seluruh pertanyaan", "Sistem menampilkan hasil analisis risiko prediabetes.", "Sistem berhasil mengirim jawaban kuesioner ke backend dan menampilkan hasil analisis risiko.", "Pass", "-", "Deny Riansyah", "30 Juli 2026"],
        ["6", "TC-006", "Riwayat Analisis", "Pengguna telah melakukan analisis", "1. Membuka menu Riwayat\n2. Memilih salah satu hasil", "Data riwayat tersedia", "Riwayat dan hasil analisis berhasil ditampilkan.", "Sistem berhasil menampilkan seluruh riwayat analisis beserta detail tingkat risiko.", "Pass", "-", "Deny Riansyah", "30 Juli 2026"],
        ["7", "TC-007", "Daftar Program 90 Hari", "Pengguna telah memperoleh hasil", "1. Membuka menu Program\n2. Melihat informasi\n3. Mulai", "Data hasil analisis", "Program Intervensi 90 Hari berhasil diaktifkan.", "Sistem berhasil menampilkan informasi Program 90 Hari dan mengaktifkannya untuk pengguna.", "Pass", "-", "Deny Riansyah", "30 Juli 2026"],
        ["8", "TC-008", "Tahapan Program", "Pengguna mengikuti Program Intervensi", "1. Membuka jadwal hari ini\n2. Membaca panduan", "Program aktif", "Pengguna berhasil mengakses modul harian.", "Sistem berhasil menampilkan materi dan panduan untuk hari tersebut.", "Pass", "-", "Deny Riansyah", "30 Juli 2026"],
        ["9", "TC-009", "Aktivitas Harian", "Pengguna mengikuti Program", "1. Membuka menu Aktivitas\n2. Mengisi aktivitas\n3. Klik Simpan", "Tidur 8 jam, Jalan 7000 langkah", "Aktivitas harian berhasil disimpan.", "Sistem berhasil menyimpan aktivitas harian pengguna dan memperbarui data log harian.", "Pass", "-", "Deny Riansyah", "30 Juli 2026"],
        ["10", "TC-010", "Streak", "Pengguna menyelesaikan aktivitas", "1. Buka Dashboard\n2. Cek indikator streak", "Aktivitas hari ini selesai", "Nilai streak bertambah.", "Sistem berhasil memperbarui nilai streak berdasarkan aktivitas harian yang selesai.", "Pass", "-", "Deny Riansyah", "30 Juli 2026"],
        ["11", "TC-011", "Progres", "Pengguna mengikuti Program", "1. Membuka menu Progres", "Data progres tersedia", "Progres program berhasil ditampilkan.", "Sistem berhasil menampilkan perkembangan Program dalam bentuk grafik.", "Pass", "-", "Deny Riansyah", "30 Juli 2026"],
        ["12", "TC-012", "Pencapaian (Badge)", "Pengguna memenuhi syarat", "1. Membuka menu Pencapaian", "Data badge tersedia", "Badge pencapaian berhasil ditampilkan.", "Sistem berhasil menampilkan daftar badge dan pencapaian yang telah diperoleh.", "Pass", "-", "Deny Riansyah", "30 Juli 2026"],
        ["13", "TC-013", "Profil", "Pengguna telah login", "1. Membuka menu Profil\n2. Ubah data\n3. Klik Simpan", "Data profil baru", "Data profil berhasil diperbarui.", "Sistem berhasil menyimpan perubahan informasi profil dan menampilkannya.", "Pass", "-", "Deny Riansyah", "30 Juli 2026"],
        ["14", "TC-014", "Logout", "Pengguna telah login", "1. Membuka menu Profil\n2. Klik Logout", "Tidak ada", "Pengguna berhasil keluar aplikasi.", "Sistem berhasil mengakhiri sesi, menghapus token otentikasi, dan kembali ke Login.", "Pass", "-", "Deny Riansyah", "30 Juli 2026"]
    ]
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Headers
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(headers):
        hdr_cells[i].text = column_name
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(8)
                
    # Data Rows
    for row in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    file_path = "Laporan_Eksekusi_Test_GluCare.docx"
    doc.save(file_path)
    print(f"SUCCESS: File tersimpan di {file_path}")

if __name__ == "__main__":
    create_word()
