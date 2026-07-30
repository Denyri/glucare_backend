import os
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

def create_word():
    doc = Document()
    
    # Menambahkan orientasi landscape agar tabel lebar muat
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = 1 # 1 is landscape
    section.page_width = new_width
    section.page_height = new_height

    doc.add_heading('#3 Defect Management', level=1)
    doc.add_paragraph('Tabel berikut mendokumentasikan daftar temuan cacat (defect) atau bug selama proses pengujian aplikasi GluCare, yang telah disesuaikan dengan ID Test Case utama.')
    
    data = [
        ["No", "ID Defect", "Tanggal Temuan", "Tester", "ID Test Case", "Priority", "Deskripsi", "Langkah Reproduksi", "Bukti Bug", "Status", "Kategori", "Assigned To", "Tanggal Fixed"],
        ["1", "DEF-001", "24-Jul-26", "Deny", "TC-002", "High", "Password salah tetap bisa login ke dalam aplikasi.", "1. Masukkan email terdaftar\n2. Masukkan password sembarangan\n3. Klik Login", "Sistem tidak memvalidasi ketidakcocokan hash password dan langsung menerbitkan token JWT, sehingga pengguna diarahkan ke Dashboard.", "Closed", "Bug / Security", "Backend Dev", "25-Jul-26"],
        ["2", "DEF-002", "25-Jul-26", "Deny", "TC-009", "Medium", "Nilai aktivitas > 24 jam (tidak logis) tetap bisa disimpan.", "1. Buka form aktivitas harian\n2. Input lama tidur 25 jam\n3. Klik Simpan", "Form log aktivitas harian berhasil menyimpan durasi tidur 25 jam di database tanpa adanya validasi batas maksimal (24 jam) di sisi server.", "Closed", "Bug / Validasi", "Backend Dev", "26-Jul-26"],
        ["3", "DEF-003", "25-Jul-26", "Deny", "TC-004", "High", "Gagal memproses Prediksi Klinis saat menerima banyak request bersamaan.", "1. Buka API Analisis Klinis\n2. Kirim request beruntun", "Terminal server Node.js terhenti dengan pesan error '429 Too Many Requests' dari API HuggingFace karena limit harian model terlampaui.", "Closed", "Bug / API", "Backend Dev", "26-Jul-26"],
        ["4", "DEF-004", "26-Jul-26", "Deny", "TC-006", "Medium", "Layar Riwayat Analisis menjadi blank (putih) jika riwayat tes masih kosong.", "1. Buat akun baru\n2. Buka menu Riwayat Analisis", "Aplikasi Flutter mengalami 'Null Pointer Exception' saat mencoba membaca array dari response JSON yang kosong, sehingga UI crash menjadi putih.", "Closed", "Bug / UI State", "Frontend Dev", "27-Jul-26"],
        ["5", "DEF-005", "27-Jul-26", "Deny", "TC-013", "Low", "Pengguna dapat menginput angka negatif pada kolom berat badan.", "1. Buka Edit Profil\n2. Ketik berat badan '-10'\n3. Simpan", "Halaman profil menampilkan indikator berat badan '-10 kg' setelah disimpan karena tidak ada fungsi validasi minimum angka nol di form Flutter.", "Closed", "Bug / UI Form", "Frontend Dev", "28-Jul-26"]
    ]
    
    table = doc.add_table(rows=1, cols=len(data[0]))
    table.style = 'Table Grid'
    
    # Populate Header
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(data[0]):
        hdr_cells[i].text = column_name
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                
    # Populate Data
    for row in data[1:]:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    file_path = "Defect_Management_GluCare_V3.docx"
    doc.save(file_path)
    print(f"SUCCESS: File tersimpan di {file_path}")

if __name__ == "__main__":
    create_word()
