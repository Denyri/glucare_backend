import os
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

def create_word():
    doc = Document()
    
    # Project Info
    doc.add_heading('Laporan Rencana Pengujian', level=1)
    doc.add_paragraph('Project Information\nNama Projek : Aplikasi GluCare\nTanggal Dokumen : 30 Juli 2026\nTim penguji : [Nama Tester/QA]\nKlien : [Company Name/Universitas]')
    
    # 1. Kebutuhan fungsional
    doc.add_heading('1. Kebutuhan fungsional', level=2)
    fr_list = [
        "FR-01: Pengguna dapat melakukan pendaftaran akun baru pada aplikasi.",
        "FR-02: Pengguna dapat melakukan login ke dalam aplikasi.",
        "FR-03: Pengguna dapat melengkapi informasi profil setelah berhasil mendaftarkan akun.",
        "FR-04: Pengguna dapat memasukkan data metrik kesehatan untuk melakukan Analisis Klinis.",
        "FR-05: Pengguna dapat mengisi kuesioner untuk melakukan Analisis Kuesioner.",
        "FR-06: Pengguna dapat melihat hasil serta riwayat analisis yang telah dilakukan, baik Analisis Klinis maupun Analisis Kuesioner.",
        "FR-07: Pengguna dapat mendaftarkan diri untuk mengikuti Program Intervensi 90 Hari.",
        "FR-08: Pengguna dapat mengikuti tahapan Program Intervensi 90 Hari sesuai dengan panduan yang diberikan aplikasi.",
        "FR-09: Pengguna dapat mencatat aktivitas harian sebagai bagian dari Program Intervensi 90 Hari.",
        "FR-10: Pengguna dapat melihat streak aktivitas harian yang menunjukkan konsistensi dalam menjalankan program.",
        "FR-11: Pengguna dapat melihat progres kesehatan selama mengikuti Program Intervensi 90 Hari.",
        "FR-12: Pengguna dapat melihat daftar pencapaian (achievement) yang berhasil diperoleh.",
        "FR-13: Pengguna dapat melihat, mengubah, dan mengelola informasi profil pengguna.",
        "FR-14: Pengguna dapat melakukan logout dari aplikasi."
    ]
    for fr in fr_list:
        p = doc.add_paragraph(style='List Number')
        p.text = fr
    
    # 2. Kebutuhan non-fungsional
    doc.add_heading('2. Kebutuhan non-fungsional', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.text = 'Performance: API backend harus mampu melayani akses hingga 50 pengguna simultan dengan response time di bawah 3 detik.'
    p = doc.add_paragraph(style='List Bullet')
    p.text = 'Usability: Antarmuka mobile harus intuitif dengan penanda warna risiko (Hijau, Kuning, Merah).'
    p = doc.add_paragraph(style='List Bullet')
    p.text = 'Security: Perlindungan JWT dan sistem pencegah double submit data (gula darah).'
    
    # 3. Rencana Pengujian Fungsional
    doc.add_heading('3. Rencana Pengujian Fungsional', level=2)
    doc.add_paragraph('Dalam menjaga kualitas dari aplikasi ini, maka direncakan beberapa skenario pengujian berdasarkan perencanaan berikut, yang mencakup pengujian dengan data valid (skenario berhasil) dan data tidak valid (skenario gagal):')
    
    test_cases = [
        # FR-01
        {
            "id": "TC-001", "desc": "Melakukan pendaftaran akun baru dengan data yang valid.",
            "use_case": "Registrasi (Berhasil)", "aktor": "Pengguna", "pre": "Pengguna berada di form Registrasi",
            "data": "Email unik, nama valid, password terisi", "scenario": "1. Isi form lengkap\n2. Klik Register", "expected": "Sistem berhasil membuat akun baru dan mengarahkan ke halaman login."
        },
        {
            "id": "TC-002", "desc": "Mencoba mendaftar menggunakan email yang sudah terdaftar di sistem.",
            "use_case": "Registrasi (Gagal)", "aktor": "Pengguna", "pre": "Pengguna berada di form Registrasi",
            "data": "Email yang sudah ada di database", "scenario": "1. Isi form dengan email bekas\n2. Klik Register", "expected": "Sistem menolak dan memunculkan error: 'Email sudah digunakan'."
        },
        
        # FR-02
        {
            "id": "TC-003", "desc": "Login ke dalam aplikasi menggunakan kombinasi email dan password terdaftar.",
            "use_case": "Login (Berhasil)", "aktor": "Pengguna", "pre": "Pengguna berada di form Login",
            "data": "Email & password valid", "scenario": "1. Input email & pass\n2. Klik Login", "expected": "Sistem memvalidasi kredensial dan mengarahkan pengguna ke Dashboard."
        },
        {
            "id": "TC-004", "desc": "Mencoba login dengan memasukkan password yang salah.",
            "use_case": "Login (Gagal)", "aktor": "Pengguna", "pre": "Pengguna berada di form Login",
            "data": "Email benar, password salah", "scenario": "1. Input email & pass salah\n2. Klik Login", "expected": "Sistem menolak askes dan memunculkan error: 'Username atau password salah'."
        },

        # FR-03
        {
            "id": "TC-005", "desc": "Menyimpan data profil pengguna secara lengkap setelah registrasi.",
            "use_case": "Lengkapi Profil (Berhasil)", "aktor": "Pengguna", "pre": "Pengguna pertama kali masuk (Onboarding)",
            "data": "Tanggal lahir, jenis kelamin", "scenario": "1. Isi kelengkapan data pribadi\n2. Klik Simpan", "expected": "Sistem berhasil memperbarui data profil di database."
        },
        {
            "id": "TC-006", "desc": "Mencoba menyimpan form profil dalam keadaan data kosong atau tidak lengkap.",
            "use_case": "Lengkapi Profil (Gagal)", "aktor": "Pengguna", "pre": "Pengguna pertama kali masuk (Onboarding)",
            "data": "Form dikosongkan", "scenario": "1. Kosongkan form\n2. Klik Simpan", "expected": "Sistem menolak penyimpanan dan memunculkan validasi form wajib diisi."
        },

        # FR-04
        {
            "id": "TC-007", "desc": "Melakukan Analisis Klinis menggunakan data kesehatan yang valid.",
            "use_case": "Analisis Klinis (Berhasil)", "aktor": "Pengguna", "pre": "Pengguna telah login dan melengkapi profil.",
            "data": "Glukosa: 95 mg/dL\nUsia: 25 Tahun\nBMI: 23", "scenario": "1. Membuka halaman Analisis Klinis.\n2. Mengisi seluruh data klinis.\n3. Menekan tombol Analisis.", "expected": "Sistem berhasil mengirim data ke backend, menerima hasil analisis dari layanan AI, dan menampilkan hasil analisis kepada pengguna."
        },
        {
            "id": "TC-008", "desc": "Melakukan Prediksi risiko dengan form metrik klinis yang tidak lengkap.",
            "use_case": "Analisis Klinis (Gagal)", "aktor": "Pengguna", "pre": "Pengguna berada di form Analisis Klinis",
            "data": "Gula darah kosong", "scenario": "1. Kosongkan beberapa field klinis\n2. Klik tombol Prediksi", "expected": "Sistem menahan request dan memunculkan peringatan harap lengkapi form."
        },

        # FR-05
        {
            "id": "TC-009", "desc": "Menjawab dan mensubmit seluruh daftar kuesioner AI dengan valid.",
            "use_case": "Analisis Kuesioner (Berhasil)", "aktor": "Pengguna", "pre": "Pengguna berada di halaman kuesioner",
            "data": "Semua 8 pertanyaan terjawab", "scenario": "1. Jawab seluruh opsi kuesioner\n2. Klik Submit", "expected": "Sistem berhasil memproses jawaban dan mengembalikan skor risiko AI."
        },
        {
            "id": "TC-010", "desc": "Melakukan submit kuesioner tanpa mengisi semua pertanyaan wajib.",
            "use_case": "Analisis Kuesioner (Gagal)", "aktor": "Pengguna", "pre": "Pengguna berada di halaman kuesioner",
            "data": "Ada pertanyaan yang terlewat", "scenario": "1. Lewati beberapa pertanyaan\n2. Klik Submit", "expected": "Sistem mematikan tombol submit atau memunculkan validasi lengkapi jawaban."
        },

        # FR-06
        {
            "id": "TC-011", "desc": "Melihat detail riwayat hasil analisis klinis dan kuesioner sebelumnya.",
            "use_case": "Riwayat Analisis (Berhasil)", "aktor": "Pengguna", "pre": "Pengguna sudah pernah melakukan tes prediksi AI",
            "data": "-", "scenario": "1. Buka menu Riwayat\n2. Pilih hasil tes sebelumnya", "expected": "Sistem menampilkan daftar riwayat lengkap beserta rekomendasi AI."
        },
        {
            "id": "TC-012", "desc": "Melihat layar riwayat analisis saat pengguna belum pernah melakukan tes sama sekali.",
            "use_case": "Riwayat Analisis (Gagal / Kosong)", "aktor": "Pengguna", "pre": "Pengguna belum pernah menggunakan fitur AI",
            "data": "-", "scenario": "1. Buka menu Riwayat", "expected": "Sistem menangani state kosong dan menampilkan pesan: 'Belum ada riwayat tes'."
        },

        # FR-07
        {
            "id": "TC-013", "desc": "Mendaftarkan diri (Enroll) ke dalam program intervensi 90 hari.",
            "use_case": "Pendaftaran Program (Berhasil)", "aktor": "Pengguna", "pre": "Pengguna belum tergabung dalam program apapun",
            "data": "Target tidur: 8 Jam\nTarget jalan kaki: 30 Menit", "scenario": "1. Masuk menu Program\n2. Tentukan target harian\n3. Klik Mulai Program 90 Hari", "expected": "Sistem menyimpan target dan mengaktifkan program dari Hari ke-1."
        },
        {
            "id": "TC-014", "desc": "Mencoba mendaftar program saat status pengguna sedang aktif di program lain.",
            "use_case": "Pendaftaran Program (Gagal)", "aktor": "Pengguna", "pre": "Pengguna sedang aktif di program 90 Hari",
            "data": "Target baru", "scenario": "1. Mencoba menekan tombol Enroll lagi", "expected": "Sistem menolak request dan memblokir pendaftaran ganda."
        },

        # FR-08
        {
            "id": "TC-015", "desc": "Membuka dan mengakses panduan serta materi aktivitas program hari ini.",
            "use_case": "Tahapan Program (Berhasil)", "aktor": "Pengguna", "pre": "Pengguna berada di Hari ke-5 program",
            "data": "Akses modul Hari ke-5", "scenario": "1. Klik detail modul/hari ini di UI", "expected": "Aplikasi menampilkan materi edukasi, panduan, dan instruksi khusus hari ini."
        },
        {
            "id": "TC-016", "desc": "Mencoba mengakses materi atau panduan program untuk hari di masa depan.",
            "use_case": "Tahapan Program (Gagal)", "aktor": "Pengguna", "pre": "Pengguna berada di Hari ke-5 program",
            "data": "Akses modul Hari ke-10", "scenario": "1. Klik modul minggu depan di daftar timeline", "expected": "Sistem menolak akses dan materi tetap terkunci (Locked state)."
        },

        # FR-09
        {
            "id": "TC-017", "desc": "Mencatat metrik harian (tidur, jalan kaki, makanan) ke dalam log program.",
            "use_case": "Catat Aktivitas Harian (Berhasil)", "aktor": "Pengguna", "pre": "Pengguna belum melaporkan aktivitas hari ini",
            "data": "Tidur: 7 jam\nLangkah: 5000", "scenario": "1. Buka menu Tracking Harian\n2. Masukkan angka log\n3. Klik Simpan", "expected": "Sistem mencatat log, progress harian terpenuhi, dan poin XP bertambah."
        },
        {
            "id": "TC-018", "desc": "Mencoba mencatat laporan gula darah untuk kedua kalinya di hari yang sama (Double Submit).",
            "use_case": "Catat Aktivitas Harian (Gagal)", "aktor": "Pengguna", "pre": "Pengguna sudah mencatat log gula darah hari ini",
            "data": "Gula darah ekstra", "scenario": "1. Buka form gula darah\n2. Submit angka baru di hari yang sama", "expected": "Sistem memblokir dan menampilkan Error 400: Anda sudah mencatat untuk hari ini."
        },

        # FR-10
        {
            "id": "TC-019", "desc": "Memeriksa perhitungan streak beruntun setelah melakukan log harian yang konsisten.",
            "use_case": "Lihat Streak (Berhasil)", "aktor": "Pengguna", "pre": "Pengguna melapor aktivitas selama beberapa hari berturut-turut",
            "data": "-", "scenario": "1. Buka Dashboard utama\n2. Cek indikator Streak", "expected": "Sistem mengakumulasikan dan menampilkan angka streak yang bernilai > 0."
        },
        {
            "id": "TC-020", "desc": "Memeriksa penyesuaian streak ketika pengguna tidak melapor (bolong) satu hari.",
            "use_case": "Lihat Streak (Gagal/Reset)", "aktor": "Pengguna", "pre": "Pengguna melewati 1 hari tanpa mengisi log harian",
            "data": "-", "scenario": "1. Buka Dashboard utama di hari berikutnya\n2. Cek indikator Streak", "expected": "Sistem mereset angka streak kembali menjadi 0 (atau memberikan penalti)."
        },

        # FR-11
        {
            "id": "TC-021", "desc": "Meninjau grafik progres metrik kesehatan yang merangkum data program.",
            "use_case": "Progres Kesehatan (Berhasil)", "aktor": "Pengguna", "pre": "Pengguna telah menjalankan program > 3 hari dan rutin mengisi data",
            "data": "Data log harian tersimpan", "scenario": "1. Buka tab Progres", "expected": "Sistem merender dan menampilkan grafik tren kesehatan berdasarkan log harian."
        },
        {
            "id": "TC-022", "desc": "Melihat tab progres kesehatan ketika data laporan harian masih kosong.",
            "use_case": "Progres Kesehatan (Gagal/Kosong)", "aktor": "Pengguna", "pre": "Pengguna baru mendaftar program (Hari 1)",
            "data": "Log harian kosong", "scenario": "1. Buka tab Progres", "expected": "Sistem mendeteksi kekurangan data dan menampilkan pesan 'Belum cukup data untuk grafik'."
        },

        # FR-12
        {
            "id": "TC-023", "desc": "Melihat status lencana (achievement) setelah syarat minimal tercapai.",
            "use_case": "Lihat Achievement (Berhasil)", "aktor": "Pengguna", "pre": "Pengguna mencapai syarat streak 7 hari",
            "data": "-", "scenario": "1. Buka menu Pencapaian/Achievement", "expected": "Sistem merubah status badge menjadi menyala dan tidak terkunci."
        },
        {
            "id": "TC-024", "desc": "Melihat status lencana (achievement) sebelum syarat minimal tercapai.",
            "use_case": "Lihat Achievement (Gagal/Terkunci)", "aktor": "Pengguna", "pre": "Pengguna baru mencapai streak 2 hari",
            "data": "-", "scenario": "1. Buka menu Pencapaian/Achievement", "expected": "Sistem menampilkan badge dalam keadaan redup / terkunci karena belum memenuhi syarat."
        },

        # FR-13
        {
            "id": "TC-025", "desc": "Melakukan perubahan data profil (contoh: berat badan) dengan format yang valid.",
            "use_case": "Kelola Profil (Berhasil)", "aktor": "Pengguna", "pre": "Pengguna berada di layar pengaturan Profil",
            "data": "Berat baru: 65 kg", "scenario": "1. Edit isian berat badan\n2. Klik Simpan Pembaruan", "expected": "Sistem berhasil memperbarui data profil pengguna ke database."
        },
        {
            "id": "TC-026", "desc": "Mencoba mengubah data profil dengan angka negatif atau format teks yang tidak diizinkan.",
            "use_case": "Kelola Profil (Gagal)", "aktor": "Pengguna", "pre": "Pengguna berada di layar pengaturan Profil",
            "data": "Berat baru: -10 kg", "scenario": "1. Ketik nilai negatif di form berat badan\n2. Klik Simpan", "expected": "Sistem menolak penyimpanan dengan pesan validasi input tidak sah."
        },

        # FR-14
        {
            "id": "TC-027", "desc": "Melakukan logout untuk keluar dari sesi (session) aplikasi saat ini.",
            "use_case": "Logout (Berhasil)", "aktor": "Pengguna", "pre": "Pengguna sedang dalam kondisi Login",
            "data": "-", "scenario": "1. Klik tombol Pengaturan\n2. Pilih Keluar/Logout", "expected": "Sistem menghapus sesi (token) di device dan mengalihkan pengguna ke layar Login."
        },
        {
            "id": "TC-028", "desc": "Mencoba memanggil fungsi atau API privat sistem dalam keadaan sudah logout.",
            "use_case": "Logout (Gagal Akses Privat)", "aktor": "Pengguna", "pre": "Pengguna sudah menekan tombol Logout",
            "data": "Panggilan API privat (contoh: /profile)", "scenario": "1. Memaksa masuk ke layar profil pasca-logout (misal via back-button)", "expected": "Sistem backend menolak akses (401 Unauthorized) dan Flutter memaksa pengguna kembali ke layar Login."
        }
    ]
    
    for i, tc in enumerate(test_cases):
        doc.add_paragraph(f'Tabel 4.{i+1} {tc["use_case"]}')
        
        table = doc.add_table(rows=9, cols=2)
        table.style = 'Table Grid'
        
        rows = table.rows
        
        # Header Row
        rows[0].cells[0].text = 'Field'
        rows[0].cells[1].text = 'Keterangan'
        for cell in rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    
        # Data Rows
        rows[1].cells[0].text = 'Test Case ID'
        rows[1].cells[1].text = tc['id']
        rows[2].cells[0].text = 'Deskripsi'
        rows[2].cells[1].text = tc['desc']
        rows[3].cells[0].text = 'Nama Use Case'
        rows[3].cells[1].text = tc['use_case']
        rows[4].cells[0].text = 'Aktor'
        rows[4].cells[1].text = tc['aktor']
        rows[5].cells[0].text = 'Pre-condition'
        rows[5].cells[1].text = tc['pre']
        rows[6].cells[0].text = 'Data Pengujian'
        rows[6].cells[1].text = tc['data']
        rows[7].cells[0].text = 'Skenario Pengujian'
        rows[7].cells[1].text = tc['scenario']
        rows[8].cells[0].text = 'Expected Output'
        rows[8].cells[1].text = tc['expected']
        
        doc.add_paragraph('')

    # 4. Rencana Pengujian Non Fungsional
    doc.add_heading('4. Rencana Pengujian Non-Fungsional (Performance)', level=2)
    doc.add_heading('4.1 Performance Testing', level=3)
    doc.add_paragraph('Tujuan dari perencanaan pengujian ini adalah memastikan sistem backend Node.js GluCare dan integrasi Model Machine Learning AI tetap responsif dan stabil saat digunakan oleh banyak pengguna secara bersamaan, khususnya saat mengakses fitur analisis prediksi maupun pelaporan Tracking harian.')
    doc.add_paragraph('Metode:\nMenggunakan skrip aplikasi Locust berbasis Python.\nFitur yang diuji performanya adalah Prediksi Klinis AI, Kuesioner, dan Input Data Tracking Harian.')
    doc.add_paragraph('Target Performa:\n- Concurrent users: 50 user\n- Ramp-up: 5 user/detik')
    doc.add_paragraph('Metrik yang Direncanakan:\nResponse Time (Aggregated / per endpoint):\n- Average (rata-rata) <= 2.8 detik\n- Median (P50) <= 2.0 detik\n- P95 (95th percentile) <= 4.5 detik\n- P99 (99th percentile) <= 6.0 detik\n- Failure rate <= 2% (total fails / total requests)')

    file_path = "Laporan_Rencana_Pengujian_GluCare_V3.docx"
    doc.save(file_path)
    print(f"SUCCESS: File tersimpan di {file_path}")

if __name__ == "__main__":
    create_word()
