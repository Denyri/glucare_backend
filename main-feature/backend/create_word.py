import os
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word():
    doc = Document()
    
    # Title
    heading = doc.add_heading('4.2.1 Pengujian Fungsional (Black Box Testing)', level=1)
    
    # Paragraph 1
    doc.add_paragraph('Pengujian fungsional dilakukan menggunakan metode Black Box Testing untuk memastikan bahwa seluruh fitur pada aplikasi GluCare telah berjalan sesuai dengan kebutuhan fungsional yang telah dirancang. Pengujian dilakukan dengan memberikan berbagai masukan (input) pada setiap fitur tanpa memperhatikan struktur kode program di baliknya, kemudian membandingkan keluaran (output) sistem dengan hasil yang diharapkan.')
    
    # Paragraph 2
    doc.add_paragraph('Berdasarkan hasil pengujian Black Box terhadap seluruh fitur utama aplikasi, seluruh fungsi berhasil berjalan sesuai dengan kebutuhan sistem dan memperoleh status Pass. Hasil tersebut menunjukkan bahwa proses autentikasi, analisis risiko prediabetes, Program Intervensi 90 Hari, pencatatan aktivitas harian, pengelolaan profil, serta integrasi dengan layanan Artificial Intelligence (AI) melalui REST API telah berfungsi dengan baik dan optimal.')
    
    # Table Title
    doc.add_heading('Tabel 4.1 Hasil Pengujian Fungsional (Black Box Testing)', level=2)
    
    # Table Data
    data = [
        ["No", "ID Test Case", "Fitur", "Skenario Pengujian", "Hasil yang Diharapkan", "Hasil Aktual", "Status"],
        ["1", "TC001", "Registrasi", "Melakukan registrasi menggunakan data (nama, email, password) yang valid.", "Akun baru berhasil dibuat dan tersimpan di sistem.", "Akun berhasil dibuat.", "Pass"],
        ["2", "TC002", "Login", "Login menggunakan email dan password yang valid.", "Sistem mengautentikasi pengguna dan mengarahkan ke halaman Dashboard.", "Sistem berhasil menampilkan halaman Dashboard.", "Pass"],
        ["3", "TC003", "Login (Negatif)", "Memasukkan password yang salah atau email yang tidak terdaftar.", "Sistem menolak akses dan menampilkan pesan kesalahan kredensial.", "Pesan kesalahan berhasil ditampilkan.", "Pass"],
        ["4", "TC004", "Lengkapi Profil", "Mengisi seluruh form data profil pengguna dengan lengkap.", "Data profil pengguna berhasil diperbarui dan disimpan ke database.", "Data berhasil disimpan.", "Pass"],
        ["5", "TC005", "Analisis Klinis", "Mengirimkan parameter data klinis (lab) yang lengkap dan valid ke AI.", "Sistem memproses data dan menampilkan hasil prediksi tingkat risiko.", "Hasil analisis AI berhasil ditampilkan.", "Pass"],
        ["6", "TC006", "Analisis Klinis (Negatif)", "Mengirimkan data klinis dengan isian yang dikosongkan atau tidak lengkap.", "Sistem menolak proses dan memunculkan notifikasi validasi pengisian.", "Validasi peringatan berhasil ditampilkan.", "Pass"],
        ["7", "TC007", "Analisis Kuesioner", "Mengisi seluruh pertanyaan kuesioner kebiasaan dan riwayat kesehatan.", "Sistem memproses jawaban dan menampilkan hasil prediksi tingkat risiko.", "Hasil analisis AI berhasil ditampilkan.", "Pass"],
        ["8", "TC008", "Program Intervensi", "Mendaftar (Enroll) ke dalam Program Intervensi Gaya Hidup 90 Hari.", "Program berhasil dimulai, sistem mulai melacak progres 90 hari.", "Program berhasil diaktifkan.", "Pass"],
        ["9", "TC009", "Aktivitas Harian", "Menyimpan catatan aktivitas harian (tidur, jalan kaki, dan pola makan).", "Data aktivitas tersimpan di database dan diakumulasikan ke progress program.", "Data berhasil tersimpan.", "Pass"],
        ["10", "TC010", "Riwayat Analisis", "Membuka halaman riwayat untuk melihat hasil prediksi AI sebelumnya.", "Sistem menampilkan seluruh daftar riwayat evaluasi pengguna.", "Riwayat evaluasi berhasil ditampilkan.", "Pass"],
        ["11", "TC011", "Progres Program", "Membuka halaman progres untuk memantau aktivitas program berjalan.", "Indikator capaian (hari, xp, streak) ditampilkan sesuai data real-time.", "Progres program berhasil ditampilkan.", "Pass"],
        ["12", "TC012", "Kelola Profil", "Mengubah data informasi akun atau mengganti kata sandi (password).", "Pembaruan informasi disetujui dan diperbarui di seluruh aplikasi.", "Data profil berhasil diperbarui.", "Pass"],
        ["13", "TC013", "Logout", "Menekan tombol keluar (logout) dari aplikasi.", "Sesi pengguna dihapus dari sistem dan dikembalikan ke halaman Login.", "Sistem berhasil kembali ke layar Login.", "Pass"],
    ]
    
    # Create Table
    table = doc.add_table(rows=1, cols=len(data[0]))
    table.style = 'Table Grid'
    
    # Populate Header
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(data[0]):
        hdr_cells[i].text = column_name
        # Simple bold for header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                
    # Populate Data
    for row in data[1:]:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    # Save to current dir
    file_path = "Pengujian_Black_Box_GluCare.docx"
    
    try:
        doc.save(file_path)
        print(f"SUCCESS: File tersimpan di {file_path}")
    except Exception as e:
        print(f"ERROR: {e}")

if __name__ == "__main__":
    create_word()
